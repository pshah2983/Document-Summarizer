import PyPDF2
import docx
import os
from transformers import pipeline
import nltk
from nltk.tokenize import sent_tokenize

nltk.download('punkt', quiet=True)

summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
qa_model = pipeline("question-answering", model="deepset/roberta-base-squad2")

def process_document(filepath):
    file_ext = os.path.splitext(filepath)[1].lower()
    if file_ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(filepath):
    doc = docx.Document(filepath)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def get_summary(text, max_length=180, min_length=60):
    # Clean and filter text for summarization
    sentences = sent_tokenize(text)
    # Remove very short/irrelevant sentences
    filtered = [s for s in sentences if len(s.split()) > 6]
    # Use only the most relevant 20 sentences (or less)
    context = " ".join(filtered[:20])
    if not context:
        context = text
    # If context is still too long, chunk it
    chunks = split_text_into_chunks(context, max_length=900)
    summaries = []
    for chunk in chunks:
        summary = summarizer(chunk, max_length=max_length, min_length=min_length, do_sample=False)
        summaries.append(summary[0]['summary_text'])
    # If multiple summaries, summarize them again
    if len(summaries) > 1:
        combined = " ".join(summaries)
        final = summarizer(combined, max_length=max_length, min_length=min_length, do_sample=False)
        return final[0]['summary_text']
    return summaries[0]

def get_answer(text, question):
    # Always use the full document context, chunked if needed
    chunks = split_text_into_chunks(text, max_length=900)
    best_score = 0
    best_answer = None
    for chunk in chunks:
        result = qa_model(question=question, context=chunk)
        if result['score'] > best_score and result['answer'].strip():
            best_score = result['score']
            best_answer = result['answer']
    if best_answer is None:
        return "I couldn't find a relevant answer in the document."
    return best_answer

def split_text_into_chunks(text, max_length=900):
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        current_length += len(word) + 1
        if current_length > max_length:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
        else:
            current_chunk.append(word)
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

def process_pdf(file_path):
    """
    Extract text from a PDF file with improved structure preservation.
    """
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            # Extract text with layout preservation
            page_text = page.extract_text()
            
            # Clean and structure the text
            page_text = clean_text(page_text)
            
            # Add page separator
            text += page_text + "\n\n"
    
    return structure_text(text)

def process_word(file_path):
    """
    Extract text from a Word document with improved structure preservation.
    """
    doc = Document(file_path)
    text = ""
    
    # Process paragraphs with their formatting
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"
    
    return structure_text(text)

def clean_text(text):
    """
    Clean and preprocess the extracted text while preserving important structure.
    """
    # Fix broken words (words with spaces in between)
    text = re.sub(r'(\w)\s+(\w)', r'\1\2', text)
    
    # Remove extra whitespace while preserving line breaks
    text = re.sub(r'\s+', ' ', text)
    
    # Preserve important separators
    text = text.replace('|', ' | ')
    text = text.replace('-', ' - ')
    
    # Fix common OCR issues
    text = re.sub(r'(\d)\s+(\d)', r'\1\2', text)  # Fix broken numbers
    text = re.sub(r'([a-zA-Z])\s+([a-zA-Z])', r'\1\2', text)  # Fix broken words
    
    # Remove multiple spaces
    text = ' '.join(text.split())
    
    return text

def structure_text(text):
    """
    Structure the text to preserve important information and relationships.
    """
    # Split into lines
    lines = text.split('\n')
    structured_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line contains key-value pairs
        if ':' in line:
            key, value = line.split(':', 1)
            structured_lines.append(f"{key.strip()}: {value.strip()}")
        else:
            # Try to fix any remaining broken words in the line
            line = re.sub(r'(\w)\s+(\w)', r'\1\2', line)
            structured_lines.append(line)
    
    return '\n'.join(structured_lines)

def chunk_text(text, chunk_size=1500):
    """
    Split text into smaller chunks while preserving context and structure.
    """
    # First, split by double newlines to preserve document structure
    sections = text.split('\n\n')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for section in sections:
        # Split section into sentences
        sentences = sent_tokenize(section)
        
        for sentence in sentences:
            # Fix any remaining broken words in the sentence
            sentence = re.sub(r'(\w)\s+(\w)', r'\1\2', sentence)
            sentence_size = len(sentence.split())
            
            if current_size + sentence_size > chunk_size:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks 